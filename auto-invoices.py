__version__ = "1.0.1"
__author__ = "Teodor Oprea"

# -------------------------------------------------| COPYRIGHT |-------------------------------------------------
# You are not allowed to use this program for personnal or business use without the written consent of the author.
# -------------------------------------------------| COPYRIGHT |-------------------------------------------------

import os
import sys
import json
import yaml
import math
import datetime
import subprocess
import pkgutil
import importlib
import re
from datetime import datetime
from pprint import pprint

if pkgutil.find_loader("pywin32") is None:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "pywin32"])

if pkgutil.find_loader("python-docx") is None:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])

if pkgutil.find_loader("art") is None:
    subprocess.check_call([sys.executable, "-m", "pip", "install", "art"])

import win32com.client

from art import *

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt, Inches

# constants

CONFIG_FILENAME = 'config.yaml'

WD_FORMAT_PDF = 17

DATE_FORMAT = '%Y-%m-%d'
DATETIME_FORMAT = '%Y-%m-%d %H:%M:%S'

DASHES = '-' * 40

IDENT_KEYS = ["first", "last", "street", "city", "province", "postal_code", "phone", "no_tps", "no_tvq", "no_membership"]
IDENT_KEYS_NAMES = ["First Name", "Last Name", "Street", "City", "Province", "Postal Code", "Phone Number", "TPS Number", "TVQ Number", "Membership Number"]

CLIENT_KEYS = ["path", "first", "last", "street", "city", "province", "postal_code", "phone"]
CLIENT_KEYS_NAMES = ["Folder Path", "First Name", "Last Name", "Street", "City", "Province", "Postal Code", "Phone Number"]

DETAILS_KEYS = ["is_health_assessment", "is_follow_up", "date", "am_pm", "from_time", "to_time", "cost_amount", "receipt_no", "payment_method"]
DETAILS_KEYS_NAMES = ["Health Assessment", "Follow Up", "Date", "AM/PM", "From Time", "To Time", "Cost Amount", "Receipt Number", "Payment Method"]

SAVEINFO_KEYS = ["path", "filename", "date", "notes"]
SAVEINFO_KEYS_NAMES = ["Saved to path", "Filename", "Date when saved", "Notes"]



def createReceipt(absFolderPath=None, ident=None, client=None, meeting_details=None):
    
    date = datetime.strptime(meeting_details['date'], DATE_FORMAT)

    if meeting_details['payment_method']==1:
        doc = Document('templates/template_1.docx')
    elif meeting_details['payment_method']==2:
        doc = Document('templates/template_2.docx')
    elif meeting_details['payment_method']==3:
        doc = Document('templates/template_3.docx')
    else:
        raise ValueError("Wrong payment method value!")

    table1, table2, table3, table4 = doc.tables[0], doc.tables[1], doc.tables[2], doc.tables[3]

    # PROVIDER IDENT
    table1.cell(1,1).text = "{} {}".format(ident['first'], ident['last'])       # therapist name
    table1.cell(2,1).text = ident['street']                                     # therapist address
    table1.cell(3,1).text = "{}, {}".format(ident['city'], ident['province'])   # therapist city
    table1.cell(3,3).text = ident['postal_code']                                # therapist postal code
    table1.cell(2,3).text = ident['phone']                                      # therapist phone

    # CLIENT IDENT
    table2.cell(1,1).text = "{} {}".format(client['first'], client['last'])      # client name
    table2.cell(2,1).text = client['street']                                     # client address
    table2.cell(3,1).text = "{}, {}".format(client['city'], client['province'])  # client city
    table2.cell(3,3).text = client['postal_code']                                # client postal code
    table2.cell(2,3).text = client['phone']                                      # client phone

    # CREATING FONT OBJECT
    obj_styles = doc.styles

    obj_font = obj_styles.add_style('Arial8', WD_STYLE_TYPE.CHARACTER).font
    obj_font.size = Pt(8)
    obj_font.name = 'Arial'

    obj_font = obj_styles.add_style('Arial9', WD_STYLE_TYPE.CHARACTER).font
    obj_font.size = Pt(9)
    obj_font.name = 'Arial'

    obj_font = obj_styles.add_style('Arial14', WD_STYLE_TYPE.CHARACTER).font
    obj_font.size = Pt(14)
    obj_font.name = 'Arial'

    obj_font = obj_styles.add_style('Calibri8', WD_STYLE_TYPE.CHARACTER).font
    obj_font.size = Pt(8)
    obj_font.name = 'Calibri'

    # SESSION INFORMATION
    if meeting_details['is_health_assessment']:
        table3.cell(3,0).paragraphs[0].add_run("X", style='Arial9').bold = True                 # health assessment
    if meeting_details['is_follow_up']:
        table3.cell(3,1).paragraphs[0].add_run("X", style='Arial9').bold = True                 # follw-up
    table3.cell(3,2).paragraphs[0].add_run(str(date.year), style='Arial8')        # year
    table3.cell(3,3).paragraphs[0].add_run(datetime.strftime(date, '%b'), style='Arial8')       # month
    table3.cell(3,4).paragraphs[0].add_run(str(date.day), style='Arial8')         # day
    table3.cell(3,5).paragraphs[0].add_run(meeting_details['from_time'], style='Arial8')        # from hour
    table3.cell(3,6).paragraphs[0].add_run(meeting_details['to_time'], style='Arial8')          # to hour
    if meeting_details['am_pm']:                                                                # PM or AM
        table3.cell(3,7).paragraphs[0].add_run("AM", style='Arial8')
    else:
        table3.cell(3,7).paragraphs[0].add_run("PM", style='Arial8')
    table3.cell(3,9).paragraphs[0].add_run(ident['street'], style='Arial9')                                 # address
    table3.cell(3,10).paragraphs[0].add_run(ident['city'], style='Arial9')                                  # city
    table3.cell(3,11).paragraphs[0].add_run(ident['province'], style='Arial9')                              # province
    table3.cell(3,12).paragraphs[0].add_run("{0:.2f} $".format(meeting_details['cost_amount']), style='Arial9')                 # amount

    table3.cell(11,12).paragraphs[0].add_run("{0:.2f} $".format(meeting_details['cost_amount']), style='Arial9').bold = True    # total amount

    table3.cell(12,9).paragraphs[0].add_run(costToWords(meeting_details['cost_amount']), style='Arial9')    # Amount in letters
    table3.cell(13,9).paragraphs[0].add_run("One (1)", style='Arial9')                                      # no of visits
    table3.cell(14,9).paragraphs[0].add_run(ident['no_tps'], style='Arial9')                                # TPS no
    table3.cell(15,9).paragraphs[0].add_run(ident['no_tvq'], style='Arial9')                                # TVQ no

    # NATUROPATH DECLARATION
    table4.cell(3,0).paragraphs[0].add_run(ident['no_membership'], style='Arial14').bold = True      # membership no

    # receipt no
    receipt_no_numbers = "".join(meeting_details['receipt_no'].split('-'))
    if len(receipt_no_numbers) != 6:
        raise ValueError("Receipt number doesnt have 6 numbers!")
    i = 1
    for char in receipt_no_numbers:
        table4.cell(3,i).paragraphs[0].add_run(char, style='Arial14')
        i+=1

    # date
    table4.cell(3,7).paragraphs[0].add_run(meeting_details['date'], style='Arial14')

    # DEFAULT FILENAMES
    filename = default_filename = "{} {} {} (#{})".format(meeting_details['date'], client['first'], client['last'], meeting_details['receipt_no'])
    docx_filename = "{}\\{}.docx".format(absFolderPath, default_filename)
    pdf_filename = "{}\\{}.pdf".format(absFolderPath, default_filename)

    # CHECK IF FILES ALREADY EXISTS
    if os.path.isfile(docx_filename) or os.path.isfile(pdf_filename):
        no_copy = 1
        while True:
            filename = "{} (Copy {})".format(default_filename, no_copy)
            docx_filename = "{}\\{}.docx".format(absFolderPath, filename)
            pdf_filename = "{}\\{}.pdf".format(absFolderPath, filename)
            if os.path.isfile(docx_filename) or os.path.isfile(pdf_filename):
                no_copy += 1
                continue
            else:
                break

    # SAVE DOCX
    doc.save(docx_filename)

    # SAVE TO PDF
    word = win32com.client.Dispatch('Word.Application')
    doc1 = word.Documents.Open(docx_filename)
    doc1.SaveAs(pdf_filename, FileFormat=WD_FORMAT_PDF)
    doc1.Close()
    word.Quit()

    return filename


def clear_screen():
    v = 0
    while v < 25:
        print("\n")
        v += 1

# returns index of answser int and 0 on back (8 answers max with backButton==True)
def choices(question, params, backButton=False):
    if (backButton and (len(params) > 8)):
        raise Exception("Too many params with backButton are given.")
    # clear_screen()
    print("\n\n{}\n{}\n{}\n".format(DASHES, question, DASHES))
    for x in range(len(params)):
        print("    ({}) {}".format(x+1,params[x]))
    if backButton:
        print("    (9) Back")
    print("\n")
    while True:
        try:
            if backButton:
                answer = int(input("Choose by typing the number (1-{} or 9): ".format(len(params))))
            else:
                answer = int(input("Choose by typing the number (1-{}): ".format(len(params))))
        except ValueError as e:
            print("Invalid answer!")
            continue
        if (1<=answer<=(len(params))):
            return answer
        elif (backButton and (answer==9)):
            return 0
        else:
            print("Invalid answer!")


def query_yes_no(question, default=None):
    """Ask a yes/no question via raw_input() and return their answer.
    "question" is a string that is presented to the user.
    "default" is the presumed answer if the user just hits <Enter>.
        It must be "yes" (the default), "no" or None (meaning
        an answer is required of the user).
    The "answer" return value is True for "yes" or False for "no".
    """
    valid = {"yes": True, "y": True, "ye": True,
             "no": False, "n": False}
    if default is None:
        prompt = " [y/n] "
    elif default == "yes":
        prompt = " [Y/n] "
    elif default == "no":
        prompt = " [y/N] "
    else:
        raise ValueError("invalid default answer: '%s'" % default)
    while True:
        sys.stdout.write(question + prompt)
        choice = input().lower()
        if default is not None and choice == '':
            return valid[default]
        elif choice in valid:
            return valid[choice]
        else:
            sys.stdout.write("Please respond with 'yes' or 'no' "
                             "(or 'y' or 'n').\n")


def numToLetter(value): #The function converts the numbers into letters. 
    if not isinstance(value, int):
        raise ValueError('This function only allows integers.')
    NUM_TO_WORD_MAPPING = {0: "", 1: "one", 2: "two", 3: "three", 4: "four", 5: "five", 6: "six", 7: "seven", 8: "eight", 9: "nine", 10: "ten", 11: "eleven", 12: "twelve",
     13: "thirteen", 20: "twenty", 30: "thirty", 50: "fifty", 80: "eighty", 10**2: "one hundred", 10**3: "one thousand", 10**5: "one hundred thousand", 10**6: "one milion"}
    if value in NUM_TO_WORD_MAPPING:
        return NUM_TO_WORD_MAPPING[value]
    elif 13<value<=19: return composeTeen(value)
    elif value>=20: return composeNumbers(value)
    else: raise ValueError('Out of range! (you excedeed 1000000)')

def composeNumbers(value):   #The function build every number biger than 40 
    if 20<=value<10**2:
        value1=int(str(value)[0])
        value2= int(str(value)[1])
        if value1==2: 
           value1='twen'
           return value1 + 'ty' + '-' + numToLetter(value2)
        elif value1==3: 
           value1='thir'
           return value1 + 'ty' + '-' + numToLetter(value2)
        elif value1==8: 
            value1='eigh'
            return value1 + 'ty' + '-' + numToLetter(value2)
        elif value1==5: 
            value1='fif'
            return value1 + 'ty' + '-' + numToLetter(value2)
        elif value % 10 == 0:
            return numToLetter(value1) + 'ty'
        else:
            return numToLetter(value1) + 'ty' + '-' + numToLetter(value2)        
    elif 10**2<=value<10**3:
        value1=int(str(value)[0])
        value2= int(str(value)[1:])
        return numToLetter(value1) + ' ' + 'hundred' + ' ' + numToLetter(value2)
    elif 10**3<=value<10**4:
        value1=int(str(value)[0])
        value2=int(str(value)[1:]) 
    elif 10**4<=value<10**5:
        value1=int(str(value)[0:2])
        value2=int(str(value)[2:])
    elif 10**5<=value<10**6:
        value1=int(str(value)[0:3])
        value2=int(str(value)[3:])
    return numToLetter(value1) + ' ' + 'thousand' + ' ' + numToLetter(value2)

def composeTeen(value): #The function takes the unit and then converts it into letter to build the word.
    value= numToLetter(int(str(value)[-1])) 
    if value=='five': value= 'fif'
    return value + 'teen'

def costToWords(cost): # ex: 100.00 => One hundred and 00/100 dollars
    if not isinstance(cost, float):
        raise ValueError('This function only allows floats.')
    if not cost >= 2.0:
        raise ValueError('The cost has to be bigger then $2.')
    dec = str(round(100 * (cost % 1)))
    if round(100 * (cost % 1)) < 10:
        dec = "0"+str(round(100 * (cost % 1)))
    return "{} and {}/100 dollars".format(numToLetter(int(math.floor(cost))), dec).capitalize()

# <<<<<<<<<<ADD VALIDATION>>>>>>>>>>
def changeIdent(filename="ident.json"): # change ident and return dictionnary
    print("\n\nYour Identification\n----------------")
    first = input("First name: ")
    last = input("Last name: ")
    street = input("Office street and street no.: ")
    city = input("Office city: ")
    province = input("Office province (ex: QC): ")
    postal_code = input("Office postal code (ex: X1X 1X1): ")
    phone = input("Office phone number (ex: 514-111-1111): ")
    no_tps = input("TPS number (ex: 00000 0000 RT 0001): ")
    no_tvq = input("TVQ number (ex: 00 0000 0000 TQ 0001): ")
    no_membership = input("Membership no. (ex: 12-3456): ")
    ident = {"first": first, "last": last, "street": street, "city": city, "province": province,
     "postal_code": postal_code, "phone": phone, "no_tps": no_tps, "no_tvq": no_tvq, "no_membership": no_membership}
    with open(filename, 'w') as fp: # saving to ident.json
        json.dump(ident, fp, indent=4)
    return ident


def deleteClient(client, filename="clients.json"):
    if os.path.isfile(filename):
        all_clients = loadData(filename)
        if client not in all_clients:
            raise
        del all_clients[all_clients.index(client)]
        with open(filename, 'w') as fp: # write to clients.json
            json.dump(all_clients, fp, indent=4) # Write data to clients.json
    else:
        raise

# add new clients to clients.json and returns new REGISTRED_CLIENTS
def registerNewClients(filename="clients.json"):

    # clients before
    if not os.path.isfile(filename):
        clients_before = None
    else:
        clients_before = loadData(filename)

    # Clients identification
    clear_screen()
    print("\n----------------")
    while True: # No of clients
        c = input("How many clients would you like to register? ")
        try:
            d = int(c)
        except ValueError:
            print("'{}' is not a valid option!".format(c))
            continue
        if d < 0:
            print("'{}' is not a valid option!".format(c))
            continue
        else:
            break

    if d==0:
        if clients_before is None:
            return None
        else:
            return clients_before

    
    registred_clients = list()
    for x in range(d): # 0,1,2,... Every client identification
        print("\n\nClient ({}/{})\n----------------".format(x+1,d))
        while True:
            while True:
                temp_folder_path = os.path.abspath(input("Client folder path: "))
                if os.path.isdir(temp_folder_path):
                    if os.access(temp_folder_path, os.W_OK):
                        break
                    else: 
                        print("The path you entered is protected")
                        continue
                else:
                    try:
                        os.makedirs(temp_folder_path)
                        break
                    except OSError:
                        print("The folder you entered cannot be created")
                        continue

            if query_yes_no("The folder path is: {}\nUse this path?".format(temp_folder_path)):
                break
            
        while True:
            temp_first = input("First name: ").capitalize()
            if re.match(r"^[-.a-zA-Z ]+$", temp_first):
                break
            else:
                print("The format is not valid")

        while True:
            temp_last = input("Last name: ").capitalize()
            if re.match(r"^[-.a-zA-Z ]+$", temp_last):
                break
            else:
                print("The format is not valid")

        temp_street = input("Street and street no.: ").capitalize()

        while True:
            temp_city = input("City: ").capitalize()
            if re.match(r"^[-.a-zA-Z ]+$", temp_city):
                break
            else:
                print("The format is not valid")

        while True:
            temp_province = input("Province (ex: QC): ").capitalize()
            if temp_province.isalpha(): break
            else:
                print("It has to contain only letters")


        zipCode = re.compile(r"^[ABCEGHJKLMNPRSTVXY]{1}\d{1}[A-Z]{1} *\d{1}[A-Z]{1}\d{1}$")
        while True:
            temp_postal_code = input("Postal code (ex: X1X 1X1): ").upper()
            if zipCode.match(temp_postal_code):
                temp_postal_code = temp_postal_code.replace(" ", "")
                if len(temp_postal_code) != 6: 
                    print("The postal code format is invalid")
                    continue
                temp_postal_code = temp_postal_code[:3] + ' ' + temp_postal_code[3:]
                break
            else:
                print("The postal code format is invalid")

        phone_re = re.compile(r"^(?:(?:\+?1\s*(?:[.-]\s*)?)?(?:\(\s*([2-9]1[02-9]|[2-9][02-8]1|[2-9][02-8][02-9])\s*\)|([2-9]1[02-9]|[2-9][02-8]1|[2-9][02-8][02-9]))\s*(?:[.-]\s*)?)?([2-9]1[02-9]|[2-9][02-9]1|[2-9][02-9]{2})\s*(?:[.-]\s*)?([0-9]{4})(?:\s*(?:#|x\.?|ext\.?|extension)\s*(\d+))?$")
        while True:
            temp_phone = input("Phone number (ex: 514-111-1111): ")
            if phone_re.match(temp_phone):
                temp_phone = temp_phone.replace(" ", "")
                break
            else:
                print("The phone number format is invalid")

        temp_client_data = {"path": temp_folder_path, "first": temp_first, "last": temp_last, "street": temp_street, "city": temp_city,
            "province": temp_province, "postal_code": temp_postal_code, "phone": temp_phone}
        registred_clients.append(temp_client_data)

    else:
        print("\nLoading...\n")
        if clients_before is not None:
            clients_before.extend(registred_clients)
            registred_clients = clients_before
        with open(filename , 'w') as fp: # write to clients.json
            json.dump(registred_clients, fp, indent=4) # Write data to clients.json
    
    return loadData(filename)


def loadData(filename):
    with open(filename, 'r') as fp:
        try:
            return json.load(fp)
        except json.decoder.JSONDecodeError:
            return None
             

def checkDictKeys(dictionnary, keys): # return None is not same size, False is not same key or True
    if (len(keys) != len(dictionnary.keys())): return None
    i = 0
    for item in dictionnary.keys():
        if (item != keys[i]): return False
        i += 1
    return True


def enterExit(msg):
    input("\n{}\n\n[PRESS ENTER TO EXIT THE PROGRAM]".format(msg))
    sys.exit()


def print_ident(ident):
    print("\n{}\n Therapist Identification:\n{}".format(DASHES, DASHES))
    for name, value in zip(IDENT_KEYS_NAMES, list(ident.values())):
        print(" * {:<20s}{}".format(name+":", value))
    print(DASHES)

def print_client(client):
    print("\n{}\n Client Identification:\n{}".format(DASHES, DASHES))
    for name, value in zip(CLIENT_KEYS_NAMES, list(client.values())):
        print(" * {:<20s}{}".format(name+":", value))
    print(DASHES)

def print_meeting_details(meeting_details):
    new_meeting_details = meeting_details.copy()
    new_meeting_details['cost_amount'] = "{0:.2f} $".format(new_meeting_details['cost_amount'])
    new_meeting_details['payment_method'] = "{} ({})".format(new_meeting_details['payment_method'],
        ["Cash/Interac", "Cheque", "Credit Card"][new_meeting_details['payment_method'] - 1])
    print("\n{}\n Meeting Details:\n{}".format(DASHES, DASHES))
    for name, value in zip(DETAILS_KEYS_NAMES, list(new_meeting_details.values())):
        print(" * {:<20s}{}".format(name+":", value))
    print(DASHES)

def print_saveinfo(saveinfo):
    print("\n{}\n Save Information:\n{}".format(DASHES, DASHES))
    for name, value in zip(SAVEINFO_KEYS_NAMES, list(saveinfo.values())):
        print(" * {:<20s}{}".format(name+":", value))
    print(DASHES)

def print_receipt(receipt):
    clear_screen()
    print_ident(receipt['IDENT'])
    print_client(receipt['CLIENT'])
    print_meeting_details(receipt['DETAILS'])
    print_saveinfo(receipt['SAVE'])
    
# Recommends a receipt number this year based on receipt_database and first_receipt_no
def getNextReceiptNo(receipt_database):
    str_year = str(datetime.now().year)
    if receipt_database is None:
        return "{}-{}".format(str(first_receipt_no).zfill(2), str_year)
    else:
        nums = []
        for receipt in receipt_database:
            splitted = receipt['DETAILS']['receipt_no'].split('-')
            if splitted[1] == str_year:
                try:
                    nums.append(int(splitted[0]))
                except Exception:
                    pass
        i = first_receipt_no
        while True:
            if i not in nums:
                return "{}-{}".format(str(i).zfill(2), str_year)
            i += 1


def askMeetingDetails(client, receipt_database):

    if receipt_database is None:
        is_first_receipt = True
    else:
        all_receipts_of_client = search_by_client_name(receipt_database, client['first'], client['last'])
        if len(all_receipts_of_client)==0:
            is_first_receipt = True
        else:
            is_first_receipt = False
            latest_receipt = findLastReceiptCreated(all_receipts_of_client)


    is_health_assessment = query_yes_no("Was it a HEALTH ASSESSMENT?")
    is_follow_up = query_yes_no("Was it a FOLLOW UP?")

    # Date
    if query_yes_no("Was the date TODAY?"):
        date = datetime.now()
        date_str = datetime.strftime(date, DATE_FORMAT)
    else:
        while True:
            date_str = input("DATE of meeting (YYYY-MM-DD) --> ")
            try:
                date = datetime.strptime(date_str, DATE_FORMAT)
                break
            except ValueError:
                print("Date is invalid or wrong format")

    while True: # from_time and am_pm
        from_time = input("FROM what hour was the meeting (hh:mm) --> ")
        try:
            am_pm = datetime.strptime(from_time,'%H:%M').strftime('%p')
        except Exception:
            print("'{}' does not have the right format!".format(from_time))
            continue
        if len(from_time.split(":")[1]) == 2:
            break
        else:
            print("'{}' does not have the right format! (minutes require 2 numbers)".format(from_time))

    while True: # to_time
        to_time = input("TO what hour was the meeting (hh:mm) --> ")
        try:
            datetime.strptime(to_time,'%H:%M')
        except Exception:
            print("'{}' does not have the right format!".format(to_time))
            continue
        if len(to_time.split(":")[1]) == 2:
            break
        else:
            print("'{}' does not have the right format! (minutes require 2 numbers)".format(to_time))

    while True: # cost
        if is_first_receipt:
            c = input("What amount did you charge with taxes (ex: 130.50) --> ")
        else:
            c = input("What amount did you charge with taxes (last time: {0:.2f}) --> ".format(latest_receipt['DETAILS']['cost_amount']))
        try:
            cost_amount = float(c)
        except ValueError:
            print("'{}' is not a number!".format(c))
            continue
        if (0 >= round(cost_amount, 2)):
            print("'{}' has to be bigger than 0!")
            continue
        else:
            break

    while True: # receipt no
        receipt_no = input("Type the receipt number (NEXT: {}) --> ".format(getNextReceiptNo(receipt_database)))
        try:
            receipt_no_year = receipt_no.split("-")[1]
        except IndexError:
            print("'{}' does not have the right format!".format(receipt_no))
            continue
        try:
            int(receipt_no.split("-")[0])
        except ValueError:
            print("'{}' does not have the right format!".format(receipt_no))
            continue
        if not ((len(receipt_no.split("-")[1])==4) and ((len(receipt_no.split("-")[0])==3) or (len(receipt_no.split("-")[0])==2))):
            print("'{}' does not have the right format!".format(receipt_no))
            continue
        try:
            receipt_no_year = int(receipt_no_year)
        except ValueError:
            print("'{}' does not have the right format!".format(receipt_no))
            continue
        if receipt_no_year != date.year:
            if not query_yes_no("Year in the receipt number is not the same. Continue anyways?"):
                continue
        exists = False
        if receipt_database is not None:
            for receipt in receipt_database:
                if receipt['DETAILS']['receipt_no'] == receipt_no:
                    exists = True
                    break
        if exists:
            print("'{}' already exists in database!".format(receipt_no))
            continue
        break

    if is_first_receipt:
        payment_method = choices("How were you paid?", ["Cash / Interac", "Cheque", "Credit card"])
    else:
        payment_method = choices("How were you paid? (last time: {})".format(latest_receipt['DETAILS']['payment_method']), 
            ["Cash / Interac", "Cheque", "Credit card"])

    return {"is_health_assessment": is_health_assessment, "is_follow_up": is_follow_up, "date": date_str, "am_pm": am_pm,
    "from_time": from_time, "to_time": to_time, "cost_amount": cost_amount, "receipt_no": receipt_no, "payment_method": payment_method}

# prompts what client it is and return client's dictionnary
def chooseClient(clients, question):
    fullname_list = list()
    for client in clients:
        fullname_list.append("{} {}".format(client['first'], client['last']))
    fullname_list.sort()
    number = choices(question, fullname_list)
    fullname = fullname_list[number-1]
    first_name = fullname.split(" ")[0]
    last_name = fullname.split(" ")[1]
    for client in clients:
        if client['first']==first_name and client['last']==last_name:
            return client
    raise Exception("Problem finding client!")

# returns absolute path to folder
def askPath(client):
    # 1. check if client['path'] is OK
    # 2. If OK ask if choose set path or choose another one
    # 3. If another one: ask absolute path and verify

    if os.path.isabs(client['path']):
        if query_yes_no("Existing path for client: {}\nUse this path?".format(client['path']), "yes"):
            if os.path.isdir(client['path']):
                if os.access(client['path'], os.W_OK):
                    return client['path']
                else:
                    print("Client path is protected; cannot write to folder.")
            else:
                try:
                    os.makedirs(client['path'])
                    return client['path']
                except Exception:
                    print("Cannot create folder")
    else:
        print("ERROR: Client path is not absolute!\nClient path ({})".format(client['path']))

    while True:
        path = os.path.abspath(input("Type path to save: "))
        if query_yes_no("Absolute path: {}\nConfirm this path?".format(path), "yes"): # confirm
            if os.path.isdir(path):
                if os.access(path, os.W_OK):
                    return path
                else:
                    print("Folder path is protected; cannot write to folder.")
            else:
                try:
                    os.makedirs(path)
                    return path
                except Exception:
                    print("Cannot create folder")
                    continue
                
# Returns receipt with particular string receipt number (None if nothing found)
def search_by_receipt_no(receipt_database, receipt_number):
    for receipt in receipt_database:
        if receipt['DETAILS']['receipt_no'] == receipt_number:
            return receipt
    return None

# Returns list of receipts with particular first and last name (empty list if nothing found)
def search_by_client_name(receipt_database, first_name, last_name):
    receipt_list = []
    for receipt in receipt_database:
        if receipt['CLIENT']['first'] == first_name and receipt['CLIENT']['last'] == last_name:
            receipt_list.append(receipt)
    return receipt_list

# With found receipts list, ask a receipt to view and return receipt
def choose_receipt(found_receipts, is_with_name):
    meeting_dates_list = []
    receipt_nums_list = []
    choice_list = []
    for receipt in found_receipts:
        meeting_dates_list.append(receipt['DETAILS']['date'])
        receipt_nums_list.append(receipt['DETAILS']['receipt_no'])
        if is_with_name:
            choice_list.append("Date: {} / Receipt #: {} / Name: {} {}".format(receipt['DETAILS']['date'], 
                receipt['DETAILS']['receipt_no'], receipt['CLIENT']['first'], receipt['CLIENT']['last']))
        else:
            choice_list.append("Date: {} / Receipt #: {}".format(receipt['DETAILS']['date'], receipt['DETAILS']['receipt_no']))
    # sort in reverse by date (sort all three lists based on dates)
    meeting_dates_list, receipt_nums_list, choice_list = (list(t) for t in zip(*sorted(zip(meeting_dates_list, receipt_nums_list, choice_list), reverse=True)))
    number = choices("\nWe found {} matching receipt(s) in the database. Choose one to view:".format(len(found_receipts)), choice_list)
    final = search_by_receipt_no(found_receipts, receipt_nums_list[number-1])
    if final is None:
        raise Exception()
    else:
        return final

# finds and returns the last created receipt in the receipt list provided
def findLastReceiptCreated(receipt_database):
    latest = ""
    latest_receipt = None
    for receipt in receipt_database:
        if receipt['SAVE']['date'] > latest:
            latest = receipt['SAVE']['date']
            latest_receipt = receipt
    if latest=="" or latest_receipt is None:
        raise Exception()
    else:
        return latest_receipt


def main(first_run=False):
    if first_run:
        # Welcome
        clear_screen()
        art = text2art("AUTO\nInvoices", font="small")
        input("{}\n\n{}\nWelcome to AUTO INVOICES v{}\nBy {}\n{}\n\n\n\n[PRESS ENTER TO START]".format(art, DASHES, __version__, __author__, DASHES))


    # Load config.yaml
    with open(CONFIG_FILENAME, 'r') as fp:
        try:
            config = yaml.safe_load(fp)
        except yaml.YAMLError:
            config = None
    
    first_receipt_no = 1 if None else config['first_receipt_no']


    # IDENT / CLIENTS / DATABASE / GENDATA

    # IDENT
    if os.path.isfile('ident.json'):
        print("\nLoading...\n")
        ident = loadData("ident.json")
        if ident is None: # not json format
            print("Identification file is not in JSON format! Please do it again.")
            ident = changeIdent()
            main()
        elif checkDictKeys(ident, IDENT_KEYS):
            print("Identification file is VALID")
        else:
            print("Identification file doesn't contain all the valid data! Please do it again.")
            ident = changeIdent()
            main()
    else:
        # Therapist identification
        ident = changeIdent()
        main()

    
    # CLIENTS loaded and verified
    if os.path.isfile('clients.json'):
        registred_clients = loadData("clients.json")
        if registred_clients is not None: # if is json format
            if len(registred_clients) != 0: # not empty
                # check if valid or not
                invalid = list()
                for one_client, index in zip(registred_clients, range(len(registred_clients))):
                    if not checkDictKeys(one_client, CLIENT_KEYS):
                        invalid.append(index)
                if len(invalid)==0:
                    path_invalid = list()
                    for one_client, index in zip(registred_clients, range(len(registred_clients))):
                        if not os.path.isabs(one_client['path']):
                            path_invalid.append(index)
                    if len(path_invalid)==0:
                        print("The registred clients file is VALID")
                    else:
                        sys.exit("Some clients does not have absolute path (indexes: {})".format(path_invalid))
                else:
                    sys.exit("Some clients in the registred clients file have invalid data format (indexes: {})".format(invalid))
            else:
                registred_clients = None
    else: # file clients.json not exists
        registred_clients = None
    
    
    # RECEIPTS_DATABASE
    if os.path.isfile('database.json'):
        receipt_database = loadData("database.json")
        if receipt_database is not None: # if is json format
            if len(receipt_database) != 0: # not empty
                # check if valid or not
                invalid = list()
                for one_receipt, index in zip(receipt_database, range(len(receipt_database))):
                    if not (checkDictKeys(one_receipt['IDENT'], IDENT_KEYS) and checkDictKeys(one_receipt['CLIENT'], CLIENT_KEYS)
                      and checkDictKeys(one_receipt['DETAILS'], DETAILS_KEYS) and checkDictKeys(one_receipt['SAVE'], SAVEINFO_KEYS)):
                        invalid.append(index)
                if len(invalid)==0:
                    # validate if absolute paths
                    path_invalid_client = list()
                    path_invalid_save = list()
                    for one_receipt, index in zip(receipt_database, range(len(receipt_database))):
                        if not os.path.isabs(one_receipt['CLIENT']['path']):            
                            path_invalid_client.append(index)
                        if one_receipt['SAVE']['path'] is not None:
                            if not os.path.isabs(one_receipt['SAVE']['path']):            
                                path_invalid_save.append(index)
                    if len(path_invalid_client)==0 and len(path_invalid_save)==0:
                        # validate if different receipt numbers
                        invalid_receipt_numbers = list()
                        invalid_receipt_numbers_indexes = list()
                        all_receipt_numbers = list()
                        seen = {}
                        for one_receipt in receipt_database:
                            all_receipt_numbers.append(one_receipt['DETAILS']['receipt_no'])
                        for x in all_receipt_numbers:
                            if x not in seen:
                                seen[x] = 1
                            else:
                                if seen[x] == 1:
                                    invalid_receipt_numbers.append(x)
                                seen[x] += 1  
                        invalid_receipt_numbers = list(set(invalid_receipt_numbers))
                        for i in invalid_receipt_numbers:
                            inv = list()
                            for receipt_number, receipt_index in zip(all_receipt_numbers, range(len(all_receipt_numbers))):
                                if receipt_number == i:
                                    inv.append(receipt_index)
                            if len(inv)!=0:
                                invalid_receipt_numbers_indexes.append(inv)
                        if len(invalid_receipt_numbers_indexes)==0:
                            print("The receipt database file is VALID")
                        else:
                            sys.exit("The receipt database file has duplicate receipt numbers.\nIndexes: {}".format(invalid_receipt_numbers_indexes))
                    else:
                        if len(path_invalid_client)!=0:
                            print("Some receipts does not have absolute paths in clients: (indexes: {})".format(path_invalid_client))
                        if len(path_invalid_save)!=0:
                            print("Some receipts does not have absolute paths in saveinfo: (indexes: {})".format(path_invalid_save))
                        sys.exit()
                else:
                    sys.exit("Some receipts in the receipt database have invalid data format (indexes: {})".format(invalid))
            else:
                receipt_database = None
    else: # file database.json not exists
        open("database.json", 'a').close()
        receipt_database = None


    # # GENDATA variable
    # if os.path.isfile('./gendata.json'):
    #     gendata = loadData("gendata.json")
    #     if gendata is not None: # if is json format
    #         # check if valid or not
    #         invalid = list()
    #         index = 0
    #         for single_text in gendata:
    #             if not checkDictKeys(single_text, ["path", "first", "last", "street", "city", "province", "postal_code", "phone"]):
    #                 invalid.append(index)
    #             index += 1
    #         if len(invalid)==0:
    #             print("Gendata file is VALID")
    #         else:
    #             sys.exit("Gendata.json has invalid data format (indexes: {})".format(invalid))
    # else: # file gendata.json not exists
    #     open("gendata.json", 'a').close()
    #     gendata = None





    # print("\nIdentification: {}".format(ident))
    # print("\nRegistred clients: {}".format(registred_clients))
    # print("\nReceipts: {}".format(receipt_database))
    # print("\nGenData: {}".format(gendata))
    

    # MAIN MENU\
    clear_screen()
    main_choice = choices("What action would you like to perform?",
     ["Create and store a new receipt",
      "Only create a receipt (without storing)",
      "Only store a receipt in database", 
      "View registred clients", 
      "Register new clients", 
      "Delete a client",
      "View your receipt database",
      "Search in your receipt database"])

    ############################################################################################
                                # Create a receipt

    if main_choice == 1 or main_choice == 2 or main_choice == 3:
        if registred_clients is None:
            print("\nYou need to register clients first!")
        else:
            clear_screen()
            if main_choice == 1:
                print("Create and store a new receipt\n{}".format(DASHES))
            elif main_choice == 2:
                print("Only create a receipt (without storing)\n{}".format(DASHES))
            elif main_choice == 3:
                print("Only store a receipt in database\n{}".format(DASHES))

            client = chooseClient(registred_clients, "For what client are you doing this receipt?")
            if main_choice == 3:
                path = None
            else:
                path = askPath(client)
            meeting_details = askMeetingDetails(client, receipt_database)                       # FIX THIS FUNCTION...
            if main_choice == 1 or main_choice == 3:
                notes = input("Add notes --> ")
            else:
                notes = None
            if main_choice == 1 or main_choice == 2:
                while True:
                    try:
                        print("\nCreating receipt...")
                        filename = createReceipt(path, ident, client, meeting_details)
                        break
                    except Exception:
                        input("\nClose all receipt related documents before retrying. <press enter to retry>")
            else:
                filename = None
            if main_choice == 1 or main_choice == 3:
                save_info = {"path": path, "filename": filename, "date": datetime.now().strftime(DATETIME_FORMAT), "notes": notes}
                receipt = {"IDENT": ident, "CLIENT": client, "DETAILS": meeting_details, "SAVE": save_info}
                with open('database.json', 'w') as fp:
                    if receipt_database is None:
                        json.dump([receipt], fp, indent=4)
                    else:
                        receipt_database.append(receipt)
                        json.dump(receipt_database, fp, indent=4)
            if main_choice == 1 or main_choice == 2:
                print("\nReceipts successfully created at:\n{}\\{}".format(path, filename))
            if main_choice == 1 or main_choice == 3:
                print("\nReceipts successfully stored in database")
        input("\n\n<press enter to return to main menu>")
        main()

    ############################################################################################
    elif main_choice == 4:          # View registred clients
        if registred_clients is None:
            print("\nYou do not have any registred clients.")
        else:
            clear_screen()
            chosen_client = chooseClient(registred_clients, "What client do you want to see?")
            clear_screen()
            print_client(chosen_client)
        input("\n<press enter to return to main menu>")
        main()
    ############################################################################################
    elif main_choice == 5:          # Register new clients
        registred_clients = registerNewClients()
        print("\nNew clients registered successfully!\n")
        input("\n<press enter to return to main menu>")
        main()
    ############################################################################################
    elif main_choice == 6:          # Delete a client
        if registred_clients is None:
            print("\nYou do not have any registred clients.")
        else:
            clear_screen()
            chosen_client = chooseClient(registred_clients, "What client do you want to DELETE?")
            if query_yes_no("Are you sure you want to permanently delete client '{} {}'?".format(chosen_client['first'], chosen_client['last']), "no"):
                deleteClient(chosen_client)
                print("'{} {}' was deleted from your registred clients list.".format(chosen_client['first'], chosen_client['last']))
        input("\n<press enter to return to main menu>")
        main()
    ############################################################################################
    elif main_choice == 7:          # View your receipt database
        if receipt_database is None:
            print("\nThe receipt database is empty.")
        else:
            clear_screen()
            search_choice = choices("What type of search would you like to do?", ["View all receipts"], True)

            if search_choice == 1:
                while True:
                    print_receipt(choose_receipt(receipt_database, True))
                    if not query_yes_no("\nDo you want to go back to the list of receipts?"):
                        break

            elif search_choice == 0:
                main()
            else:
                raise Exception("An unknown choice was given!")
        input("\n<press enter to return to main menu>")
        main()


    ############################################################################################
    elif main_choice == 8:          # Search in your receipt database

        if receipt_database is None:
            print("\nThe receipt database is empty.")
        else:
            clear_screen()
            search_choice = choices("What type of search would you like to do?", ["Search lastest saved receipt", "Search by receipt number", "Search by client", "Search by date"], True)
            
            # Latest saved
            if search_choice == 1:
                print_receipt(findLastReceiptCreated(receipt_database))

            # search by receipt number
            elif search_choice == 2:
                num = input("\nEnter the receipt number --> ")
                found_receipt = search_by_receipt_no(receipt_database, num)
                if found_receipt is None:
                    print("\n'{}' does not exist in the receipt database".format(num))
                else:
                    print_receipt(found_receipt)

            # search by client
            elif search_choice == 3:
                clear_screen()
                client = chooseClient(registred_clients, "For what client are you searching the receipts?")
                found_receipts = search_by_client_name(receipt_database, client['first'], client['last'])
                if len(found_receipts)==0:
                    print("No receipt found with the name {} {}.".format(client['first'], client['last']))
                else:
                    while True:
                        print_receipt(choose_receipt(found_receipts, False))
                        if not query_yes_no("\nDo you want to go back to your search results?"):
                            break

            # search by date
            elif search_choice == 4:
                while True:
                    date_str = input("\nEnter date (YYYY-MM-DD) --> ")
                    try:
                        datetime.strptime(date_str, DATE_FORMAT)
                        break
                    except ValueError:
                        print("Date is invalid or wrong format")
                found_receipts = []
                for receipt in receipt_database:
                    if receipt['DETAILS']['date'] == date_str:
                        found_receipts.append(receipt)
                if len(found_receipts)==0:
                    print("No receipt found with at the date of {}".format(date_str))
                else:
                    while True:
                        print_receipt(choose_receipt(found_receipts, True))
                        if not query_yes_no("\nDo you want to go back to your search results?"):
                            break










            elif search_choice == 0:
                main()
            else:
                raise Exception("An unknown choice was given!")
        input("\n<press enter to return to main menu>")
        main()
    ############################################################################################
    else:
        raise Exception("An unknown choice was given!")




if __name__ == '__main__':
    try:
        main(True)
    except EOFError:
        sys.exit()